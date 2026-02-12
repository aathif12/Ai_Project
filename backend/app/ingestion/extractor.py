"""
Document text extraction for PDF and DOCX files.
"""

import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)


class ExtractedPage:
    """Represents extracted content from a single page."""
    
    def __init__(self, page_number: int, text: str, metadata: Dict[str, Any] = None):
        self.page_number = page_number
        self.text = text
        self.metadata = metadata or {}


class DocumentExtractor:
    """Extract text from PDF and DOCX documents."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
    
    @classmethod
    def extract(cls, file_path: str) -> List[ExtractedPage]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of ExtractedPage objects with text content
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {cls.SUPPORTED_EXTENSIONS}"
            )
        
        if extension == ".pdf":
            return cls._extract_pdf(file_path)
        elif extension == ".docx":
            return cls._extract_docx(file_path)
    
    @classmethod
    def _extract_pdf(cls, file_path: str) -> List[ExtractedPage]:
        """Extract text from PDF using PyMuPDF."""
        pages = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                
                # Clean up the text
                text = cls._clean_text(text)
                
                if text.strip():  # Only add non-empty pages
                    pages.append(ExtractedPage(
                        page_number=page_num + 1,  # 1-indexed
                        text=text,
                        metadata={
                            "source": file_path,
                            "page_width": page.rect.width,
                            "page_height": page.rect.height,
                        }
                    ))
            
            doc.close()
            logger.info(f"Extracted {len(pages)} pages from PDF: {file_path}")
            
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise
        
        return pages
    
    @classmethod
    def _extract_docx(cls, file_path: str) -> List[ExtractedPage]:
        """
        Extract text from DOCX.
        Note: DOCX doesn't have traditional pages, so we treat the whole doc as one page.
        """
        pages = []
        
        try:
            doc = Document(file_path)
            
            # Extract all paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            combined_text = "\n\n".join(full_text)
            combined_text = cls._clean_text(combined_text)
            
            if combined_text.strip():
                pages.append(ExtractedPage(
                    page_number=1,
                    text=combined_text,
                    metadata={
                        "source": file_path,
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                    }
                ))
            
            logger.info(f"Extracted text from DOCX: {file_path}")
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
        
        return pages
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text by removing excess whitespace."""
        # Replace multiple newlines with double newline
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    @classmethod
    def get_page_count(cls, file_path: str) -> int:
        """Get the number of pages in a document."""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == ".pdf":
            doc = fitz.open(file_path)
            count = len(doc)
            doc.close()
            return count
        elif extension == ".docx":
            return 1  # DOCX doesn't have traditional pages
        
        return 0
